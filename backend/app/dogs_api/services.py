import os
import subprocess
from dataclasses import dataclass

from django.conf import settings
from django.core.mail import EmailMessage
from docx import Document

from app.dogs_api.domain import (
    DocxFillError,
    EmailSendError,
    FormularioAdopcion,
    PdfConversionError,
)


@dataclass
class AdoptionFormManager:
    domain_object: FormularioAdopcion
    TEMPLATE_DOCX_PATH: str = "app/form_templates/Prueba_formulario_1.docx"
    FILLED_DOCX_PATH: str = "app/form_templates/filled_form.docx"
    FILLED_PDF_DIR: str = "app/form_templates/"
    FILLED_PDF_PATH: str = "app/form_templates/filled_form.pdf"

    def execute(self):
        try:
            self.fill_docx()
            self.docx_to_pdf()
            self.send_mail()
        finally:
            self.cleanup()

    def fill_docx(self):
        try:
            doc = Document(self.TEMPLATE_DOCX_PATH)
            context = self.domain_object.make_data_dict()

            for paragraph in doc.paragraphs:
                if not any(f"{{{{{k}}}}}" in paragraph.text for k in context):
                    continue

                runs = paragraph.runs
                full_text = "".join(run.text for run in runs)

                for key, val in context.items():
                    placeholder = f"{{{{{key}}}}}"
                    if isinstance(val, bool):
                        val = "Si" if val else "No"
                    elif isinstance(val, (int, float)):
                        val = str(val)
                    full_text = full_text.replace(placeholder, val)

                if runs:
                    runs[0].text = full_text
                    for r in runs[1:]:
                        r.text = ""
                else:
                    paragraph.add_run(full_text)

            doc.save(self.FILLED_DOCX_PATH)
        except FileNotFoundError as e:
            raise DocxFillError(f"Error: El archivo plantilla no fue encontrado en '{self.TEMPLATE_DOCX_PATH}'.") from e
        except Exception as e:
            raise DocxFillError(f"Ocurrió un error inesperado al rellenar el documento: {e}") from e

    def docx_to_pdf(self):
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", self.FILLED_DOCX_PATH, "--outdir", self.FILLED_PDF_DIR],
                check=True,
                capture_output=True,  # Captura stdout y stderr
            )
        except FileNotFoundError as e:
            raise PdfConversionError("Error: 'libreoffice' no está instalado o no se encuentra en el PATH del sistema.") from e
        except subprocess.CalledProcessError as e:
            error_message = e.stderr.decode('utf-8') if e.stderr else "No se pudo obtener el error específico de LibreOffice."
            raise PdfConversionError(f"LibreOffice falló al convertir el archivo a PDF: {error_message}") from e
        except Exception as e:
            raise PdfConversionError(f"Ocurrió un error inesperado durante la conversión a PDF: {e}") from e

    def send_mail(self):
        html = f"<h1>Solicitud de adopcion recibida!!!</h1><p>Una persona ha solicitado la adopcion de {self.domain_object.datos_del_animal.dog_name}, los datos de la adopcion se encuentran en el siguiente pdf</p>"
        email = EmailMessage(
            subject="Solicitud de adopcion Esperanza Canina",
            body=html,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=settings.DESTINATARY_EMAIL,
        )
        email.content_subtype = "html"
        email.attach_file(self.FILLED_PDF_PATH)
        try:
            email.send(fail_silently=False)
            print("Correo HTML enviado exitosamente con EmailMessage.")
        except Exception as e:
            # Idealmente, capturar excepciones más específicas de SMTP si es posible.
            raise EmailSendError(f"Error al enviar correo: {e}") from e

    def cleanup(self):
        for path in [self.FILLED_DOCX_PATH, self.FILLED_PDF_PATH]:
            try:
                if os.path.exists(path):
                    os.remove(path)
                    print(f"Archivo temporal eliminado: {path}")
            except OSError as e:
                # Loggear este error en lugar de solo imprimirlo sería una buena mejora
                print(f"Error al eliminar el archivo {path}: {e}")
