import subprocess
import time

from docx import Document

from app.dogs_api.adoption_form_domain import make_example_form


def fill_docx(template_path: str, output_path: str, context: dict):
    """
    Rellena placeholders tipo {{key}} en los párrafos de un .docx y guarda el resultado.
    - template_path: ruta al .docx plantilla
    - output_path: ruta donde se guardará el .docx resultante
    - context: dict con pares { "key": "valor" }
    """
    doc = Document(template_path)

    # Recorremos todos los párrafos del documento
    for paragraph in doc.paragraphs:
        # Si el párrafo no contiene ninguno de los placeholders, lo saltamos rápido
        if not any(f"{{{{{k}}}}}" in paragraph.text for k in context):
            continue

        # Los párrafos están compuestos por "runs" (fragmentos con formato propio).
        # Muchas veces un placeholder se puede partir en varios runs, así que:
        runs = paragraph.runs
        full_text = "".join(run.text for run in runs)

        # Reemplazamos todos los placeholders en el texto completo del párrafo
        for key, val in context.items():
            placeholder = f"{{{{{key}}}}}"
            if isinstance(val, bool):
                val = "Si" if val else "No"
            elif isinstance(val, (int, float)):
                val = str(val)
            full_text = full_text.replace(placeholder, val)

        # Escribimos el texto resultante de vuelta en los runs:
        # - dejamos el primer run con todo el texto nuevo (preserva su estilo),
        # - vaciamos los runs siguientes para evitar texto duplicado.
        if runs:
            runs[0].text = full_text
            for r in runs[1:]:
                r.text = ""
        else:
            # Si por alguna razón no hay runs (raro), creamos uno nuevo
            paragraph.add_run(full_text)

    # Guardar documento modificado
    doc.save(output_path)



def docx_to_pdf(input_path, output_path):
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", output_path
    ], check=True)

if __name__ == "__main__":
    start = time.time()
    domain = make_example_form()
    data = domain.make_data_dict()

    fill_docx("app/form_templates/Prueba_formulario_1.docx", "app/form_templates/filled_form.docx", data)
    docx_to_pdf("app/form_templates/filled_form.docx", "app/form_templates/")
    end = time.time()
    print(f"Tiempo de ejecución: {end - start:.2f} segundos")
